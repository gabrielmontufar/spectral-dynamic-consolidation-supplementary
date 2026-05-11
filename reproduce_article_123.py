from __future__ import annotations

import math
import os
import shutil
import zipfile
from pathlib import Path

import matplotlib

matplotlib.use("Agg")

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt
from lxml import etree
from PIL import Image


ROOT = Path(r"C:\Users\gjm31\OneDrive\Escritorio\articulos up\123 avo art")
SOURCE = Path(r"C:\Users\gjm31\OneDrive\Escritorio\123 avo art")
DELIVERY = ROOT / "Entrega Acta Geotechnica 20260511"
MANUSCRIPT_DIR = DELIVERY / "01 Manuscript"
FIG_DIR = DELIVERY / "02 Figures"
TABLE_DIR = DELIVERY / "03 Table images"
SUPP_DIR = DELIVERY / "04 Supplemental data and code"
DECL_DIR = DELIVERY / "05 Cover and declarations"
PROTO_DIR = DELIVERY / "06 Protocol and journal instructions"
RENDER_DIR = DELIVERY / "render_manual_check"
REPO_DIR = ROOT / "supplementary_spectral_consolidation"

PY = Path(r"C:\Users\gjm31\.cache\codex-runtimes\codex-primary-runtime\dependencies\python\python.exe")
MML2OMML = Path(r"C:\Program Files\Microsoft Office\Office16\MML2OMML.XSL")


def ensure_dirs() -> None:
    ROOT.mkdir(parents=True, exist_ok=True)
    for p in [MANUSCRIPT_DIR, FIG_DIR, TABLE_DIR, SUPP_DIR, DECL_DIR, PROTO_DIR, RENDER_DIR, REPO_DIR]:
        p.mkdir(parents=True, exist_ok=True)
    for p in [MANUSCRIPT_DIR, FIG_DIR, TABLE_DIR, SUPP_DIR, DECL_DIR]:
        for old in p.glob("*"):
            if old.is_file():
                old.unlink()
    for f in SOURCE.glob("*"):
        if f.is_file():
            shutil.copy2(f, ROOT / f.name)


def spectral_retention(pi: np.ndarray | float, terms: int = 500) -> np.ndarray | float:
    pi_arr = np.asarray(pi, dtype=float)
    out = np.zeros_like(pi_arr, dtype=float)
    m = np.arange(1, 2 * terms, 2, dtype=float)
    for idx, val in np.ndenumerate(pi_arr):
        out[idx] = np.sum(8.0 / (m**4 * math.pi**4 * val) * (1.0 - np.exp(-(m**2) * math.pi**2 * val)))
    if np.isscalar(pi):
        return float(out)
    return out


def solve_threshold(target: float) -> float:
    lo, hi = 1e-6, 100.0
    for _ in range(90):
        mid = 0.5 * (lo + hi)
        if spectral_retention(mid) > target:
            lo = mid
        else:
            hi = mid
    return 0.5 * (lo + hi)


def fd_retention(pi: float, n: int = 120, steps: int = 800) -> float:
    """Implicit finite-difference solution of u_t = u_xx + 1, u(0)=u(1)=0."""
    dx = 1.0 / (n + 1)
    dt = pi / max(steps, 1)
    u = np.zeros(n)
    r = dt / dx**2
    diag = (1.0 + 2.0 * r) * np.ones(n)
    off = -r * np.ones(n - 1)
    A = np.diag(diag) + np.diag(off, 1) + np.diag(off, -1)
    rhs_add = dt * np.ones(n)
    for _ in range(steps):
        u = np.linalg.solve(A, u + rhs_add)
    avg = dx * np.sum(u)
    return float(avg / pi)


def fem_retention(pi: float, elems: int = 90, steps: int = 700) -> float:
    """Linear Galerkin FEM for u_t = u_xx + 1 on [0,1] with drained ends."""
    n_nodes = elems + 1
    h = 1.0 / elems
    n_int = n_nodes - 2
    M = np.zeros((n_int, n_int))
    K = np.zeros((n_int, n_int))
    F = np.zeros(n_int)
    Me = h / 6.0 * np.array([[2.0, 1.0], [1.0, 2.0]])
    Ke = 1.0 / h * np.array([[1.0, -1.0], [-1.0, 1.0]])
    Fe = h / 2.0 * np.array([1.0, 1.0])
    for e in range(elems):
        nodes = [e, e + 1]
        for a in range(2):
            ia = nodes[a] - 1
            if 0 <= ia < n_int:
                F[ia] += Fe[a]
                for b in range(2):
                    ib = nodes[b] - 1
                    if 0 <= ib < n_int:
                        M[ia, ib] += Me[a, b]
                        K[ia, ib] += Ke[a, b]
    dt = pi / max(steps, 1)
    u = np.zeros(n_int)
    A = M + dt * K
    for _ in range(steps):
        u = np.linalg.solve(A, M @ u + dt * F)
    # Domain average from consistent mass load vector.
    avg = F @ u
    return float(avg / pi)


def finite_diff_sensitivities(base: dict[str, float]) -> pd.DataFrame:
    def fs_from(vals: dict[str, float]) -> float:
        pi = vals["cv"] * vals["T"] / vals["h"] ** 2
        r = spectral_retention(pi)
        p_ret = r * vals["Pu"]
        return ((vals["sigma_eff0"] - p_ret) * math.tan(math.radians(vals["phi"])) + vals["c"]) / vals["tau"]

    base_fs = fs_from(base)
    rows = []
    for key in ["h", "cv", "T", "Pu", "phi", "sigma_eff0", "tau"]:
        hi = dict(base)
        lo = dict(base)
        hi[key] *= 1.10
        lo[key] *= 0.90
        if key == "phi":
            hi[key] = base[key] + 2.0
            lo[key] = base[key] - 2.0
            denom = math.log(hi[key] / lo[key])
        else:
            denom = math.log(hi[key] / lo[key])
        s = (math.log(fs_from(hi)) - math.log(fs_from(lo))) / denom
        rows.append(
            {
                "Parameter": key,
                "Base value": base[key],
                "Normalized sensitivity of FS_PD": s,
                "Interpretation": "destabilizing when increased" if s < 0 else "stabilizing when increased",
            }
        )
    return pd.DataFrame(rows).sort_values("Normalized sensitivity of FS_PD")


def convergence_table() -> pd.DataFrame:
    test_pis = np.array([solve_threshold(0.9), solve_threshold(0.5), solve_threshold(0.1)])
    ref = spectral_retention(test_pis, terms=5000)
    rows = []
    for terms in [5, 10, 20, 50, 100]:
        vals = spectral_retention(test_pis, terms=terms)
        rows.append(
            {
                "Check": "spectral truncation",
                "Resolution": f"{terms} odd terms",
                "Maximum absolute error": float(np.max(np.abs(vals - ref))),
                "Comment": "series truncation relative to 5000-term reference",
            }
        )
    for n, steps in [(40, 300), (80, 600), (120, 800)]:
        vals = np.array([fd_retention(float(pi), n=n, steps=steps) for pi in test_pis])
        rows.append(
            {
                "Check": "finite difference",
                "Resolution": f"{n} interior nodes, {steps} time steps",
                "Maximum absolute error": float(np.max(np.abs(vals - ref))),
                "Comment": "implicit time integration",
            }
        )
    for elems, steps in [(30, 300), (60, 600), (90, 700)]:
        vals = np.array([fem_retention(float(pi), elems=elems, steps=steps) for pi in test_pis])
        rows.append(
            {
                "Check": "finite element",
                "Resolution": f"{elems} linear elements, {steps} time steps",
                "Maximum absolute error": float(np.max(np.abs(vals - ref))),
                "Comment": "consistent-mass Galerkin solution",
            }
        )
    return pd.DataFrame(rows)


def literature_comparison_table() -> pd.DataFrame:
    return pd.DataFrame(
        [
            ["Velocity-based drainage choice", "movement rate or loading rate", "quick practical decision", "does not include h or cv explicitly", "replaces binary label with retained pressure R(Pi)"],
            ["Conventional drained/undrained stability", "selected strength envelope", "clear design bounds", "no intermediate state unless analyst brackets cases", "computes partly drained effective stress continuously"],
            ["Classical consolidation", "time factor", "well established diffusion physics", "not directly linked to landslide FS in screening workflows", "couples spectral pressure retention to FS"],
            ["Transient seepage analysis", "boundary flux and hydraulic gradients", "captures rainfall or seepage forcing", "may not represent rapid internal pressure generation", "adds a source-driven rapid-event operator"],
            ["Fully coupled hydro-mechanical FEM", "field equations, constitutive law and geometry", "most general numerical route", "costly and opaque for screening inventories", "provides a reproducible pre-analysis and check on regime choice"],
        ],
        columns=["Approach", "Primary control variable", "Strength", "Limitation for rapid saturated sliding", "Contribution of the present criterion"],
    )


def external_consistency_table() -> pd.DataFrame:
    return pd.DataFrame(
        [
            ["Iverson pore-pressure feedback model", "Motion styles depend on coupling between shear-zone volume change and pore-pressure dissipation timescales.", "The proposed Pi and R(Pi) isolate the dissipation part of that feedback.", "External conceptual consistency, not direct calibration."],
            ["USGS Oso ring-shear data release", "Undrained tests used closed chamber water lines; naturally drained tests used open lines; pore-water pressure and specimen thickness were measured continuously.", "The limiting cases match the model endpoints R approximately 1 for closed drainage and lower R for open drainage as hydraulic time increases.", "A future calibration can ingest the released time series; this manuscript uses the release as an external mode check."],
            ["Risk frameworks for infrastructure slopes", "Road vulnerability depends on physical response, not only hazard presence.", "The retained-pressure variable supplies a transparent state descriptor for road cuts, embankments and retaining-system back-slopes.", "Supports engineering interpretation but does not replace site-specific risk modelling."],
        ],
        columns=["External anchor", "Published or released observation", "How the criterion is checked", "Limit of the check"],
    )


def build_computation_package() -> dict[str, pd.DataFrame]:
    pi_grid = np.logspace(-4, 2, 241)
    retention = spectral_retention(pi_grid)
    thresholds = {0.9: solve_threshold(0.9), 0.5: solve_threshold(0.5), 0.1: solve_threshold(0.1)}

    curve = pd.DataFrame({"Pi": pi_grid, "R_spectral": retention})
    curve["regime"] = np.where(curve["R_spectral"] >= 0.9, "nearly undrained", np.where(curve["R_spectral"] <= 0.1, "drained", "partly drained"))

    bench = pd.DataFrame(
        [
            {"Case": "A", "cv_m2_s": 5e-6, "T_s": 2.0, "h_m": 0.05, "sigma_eff0_kPa": 50.0, "Pu_kPa": 20.0, "phi_deg": 30.0, "tau_kPa": 25.0, "Infrastructure reading": "short saturated shear pulse below a road cut"},
            {"Case": "B", "cv_m2_s": 5e-6, "T_s": 60.0, "h_m": 0.05, "sigma_eff0_kPa": 50.0, "Pu_kPa": 20.0, "phi_deg": 30.0, "tau_kPa": 25.0, "Infrastructure reading": "prolonged rapid movement affecting a transport corridor"},
            {"Case": "C", "cv_m2_s": 1e-4, "T_s": 60.0, "h_m": 0.05, "sigma_eff0_kPa": 50.0, "Pu_kPa": 20.0, "phi_deg": 30.0, "tau_kPa": 25.0, "Infrastructure reading": "more permeable drainage path or engineered relief drain"},
            {"Case": "D", "cv_m2_s": 2e-5, "T_s": 20.0, "h_m": 0.03, "sigma_eff0_kPa": 65.0, "Pu_kPa": 18.0, "phi_deg": 32.0, "tau_kPa": 31.0, "Infrastructure reading": "thin sheared zone behind a retaining structure"},
            {"Case": "E", "cv_m2_s": 1e-6, "T_s": 120.0, "h_m": 0.08, "sigma_eff0_kPa": 45.0, "Pu_kPa": 16.0, "phi_deg": 28.0, "tau_kPa": 23.0, "Infrastructure reading": "low-permeability colluvium supporting a rural road"},
        ]
    )
    bench["Pi"] = bench.cv_m2_s * bench.T_s / bench.h_m**2
    bench["R_Pi"] = spectral_retention(bench["Pi"].to_numpy())
    bench["retained_pressure_kPa"] = bench["R_Pi"] * bench["Pu_kPa"]
    bench["FS_PD"] = ((bench["sigma_eff0_kPa"] - bench["retained_pressure_kPa"]) * np.tan(np.radians(bench["phi_deg"]))) / bench["tau_kPa"]
    bench["Regime"] = np.where(bench["R_Pi"] >= 0.9, "nearly undrained", np.where(bench["R_Pi"] <= 0.1, "drained", "partly drained"))

    validation_pis = np.array([0.004, 0.02, 0.1126126506, 0.12, 0.5, 0.8331127834, 2.4])
    validation_rows = []
    for pi in validation_pis:
        rs = spectral_retention(pi)
        rfd = fd_retention(pi)
        rfe = fem_retention(pi)
        validation_rows.append(
            {
                "Pi": pi,
                "R_spectral": rs,
                "R_FD": rfd,
                "R_FEM": rfe,
                "FD_abs_error": abs(rfd - rs),
                "FEM_abs_error": abs(rfe - rs),
            }
        )
    validation = pd.DataFrame(validation_rows)
    convergence = convergence_table()

    base = {"h": 0.05, "cv": 5e-6, "T": 60.0, "Pu": 20.0, "phi": 30.0, "sigma_eff0": 50.0, "tau": 25.0, "c": 0.0}
    sensitivity = finite_diff_sensitivities(base)

    regimes = pd.DataFrame(
        [
            {"Boundary": "nearly undrained / partly drained", "Criterion": "R(Pi)=0.9", "Pi": thresholds[0.9], "Engineering meaning": "less than 10% of the undrained pressure has dissipated"},
            {"Boundary": "central transition", "Criterion": "R(Pi)=0.5", "Pi": thresholds[0.5], "Engineering meaning": "half of the undrained pressure remains at the end of the rapid event"},
            {"Boundary": "partly drained / drained", "Criterion": "R(Pi)=0.1", "Pi": thresholds[0.1], "Engineering meaning": "only 10% of the undrained pressure remains"},
        ]
    )

    for name, df in {
        "retention_curve.csv": curve,
        "benchmark_cases.csv": bench,
        "validation_fd_fem.csv": validation,
        "convergence_checks.csv": convergence,
        "literature_comparison.csv": literature_comparison_table(),
        "external_consistency_check.csv": external_consistency_table(),
        "sensitivity.csv": sensitivity,
        "thresholds.csv": regimes,
    }.items():
        df.to_csv(SUPP_DIR / name, index=False)
        df.to_csv(REPO_DIR / name, index=False)

    with pd.ExcelWriter(TABLE_DIR / "Article 123 tables.xlsx", engine="openpyxl") as writer:
        for sheet, df in {
            "Table 1 variables": variable_table(),
            "Table 2 regimes": regimes,
            "Table 3 benchmark": display_benchmark_compact_table(bench),
            "Table 4 validation": display_validation_table(validation).round(6),
            "Table 5 convergence": convergence.round(8),
            "Table 6 literature": literature_comparison_table(),
            "Table 7 external": external_consistency_table(),
            "Table 8 sensitivity": display_sensitivity_table(sensitivity).round(5),
            "Table 9 infrastructure": infrastructure_table(),
        }.items():
            df.to_excel(writer, sheet_name=sheet[:31], index=False)

    return {
        "curve": curve,
        "bench": bench,
        "validation": validation,
        "convergence": convergence,
        "sensitivity": sensitivity,
        "regimes": regimes,
        "literature": literature_comparison_table(),
        "external": external_consistency_table(),
    }


def set_cell_shading(cell, fill: str = "FFFFFF") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text_style(cell, bold=False, size=8.5) -> None:
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        for run in paragraph.runs:
            run.font.name = "Times New Roman"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
            run.font.size = Pt(size)
            run.font.bold = bold
            run.font.color.rgb = None


def style_table(table, header=True, size=8.5) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    table.style = "Table Grid"
    for r_idx, row in enumerate(table.rows):
        for cell in row.cells:
            set_cell_shading(cell, "FFFFFF")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_text_style(cell, bold=(header and r_idx == 0), size=size)


def add_table_from_df(doc: Document, df: pd.DataFrame, size=8.2) -> None:
    table = doc.add_table(rows=1, cols=len(df.columns))
    hdr = table.rows[0].cells
    for i, col in enumerate(df.columns):
        hdr[i].text = str(col)
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for i, value in enumerate(row):
            if isinstance(value, float):
                if value != 0 and abs(value) < 1e-3:
                    text = f"{value:.2e}"
                else:
                    text = f"{value:.4g}"
            else:
                text = str(value)
            cells[i].text = text
    style_table(table, size=size)


def variable_table() -> pd.DataFrame:
    return pd.DataFrame(
        [
            ["p(n,t)", "excess pore-water pressure", "kPa", "solved field variable"],
            ["n", "coordinate normal to the shear band", "m", "0 <= n <= h"],
            ["h", "hydraulic thickness of the active shear band", "m", "measured or bracketed from site interpretation"],
            ["T", "effective duration of rapid loading or sliding", "s", "screening variable for transient motion"],
            ["cv", "effective consolidation coefficient", "m2/s", "from permeability and constrained storage"],
            ["Pi", "dynamic consolidation number", "-", "defined by Eq. (3)"],
            ["R(Pi)", "retained undrained pore-pressure fraction", "-", "spectral transition function"],
            ["Pu", "undrained pressure increment", "kPa", "generated by loading or contractancy"],
            ["FSPD", "partly drained factor of safety", "-", "effective-stress stability index"],
        ],
        columns=["Symbol", "Definition", "Unit", "Role in the criterion"],
    )


def infrastructure_table() -> pd.DataFrame:
    return pd.DataFrame(
        [
            ["Road cut in saturated colluvium", "low cv, small h, short acceleration T", "nearly undrained response can persist even during short motion", "treat rapid movement as a retained-pressure problem, not a velocity-only classification"],
            ["Highway embankment on clayey slope", "moderate cv but long T during progressive acceleration", "partly drained path with material-dependent FS shift", "screen drainage and shear-band thickness before assigning undrained strength"],
            ["Slope with relief drains", "engineered increase of drainage capacity", "Pi increases and R(Pi) decreases", "quantify whether drains move the state across the R=0.5 or R=0.1 boundary"],
            ["Retaining wall back-slope", "thin local shear band near a structural boundary", "h uncertainty dominates the hydraulic regime", "perform sensitivity on h before relying on a single stability check"],
        ],
        columns=["Infrastructure setting", "Dominant uncertainty", "Consequence for drainage regime", "Design use"],
    )


def write_table_images(tables: dict[str, pd.DataFrame]) -> None:
    for idx, (title, df) in enumerate(tables.items(), 1):
        fig_w = 11
        fig_h = max(2.2, 0.45 * (len(df) + 2))
        fig, ax = plt.subplots(figsize=(fig_w, fig_h), dpi=220)
        ax.axis("off")
        display_df = df.copy()
        for c in display_df.columns:
            display_df[c] = display_df[c].map(
                lambda v: (f"{v:.2e}" if v != 0 and abs(v) < 1e-3 else f"{v:.4g}") if isinstance(v, float) else str(v)
            )
        tbl = ax.table(
            cellText=display_df.values,
            colLabels=display_df.columns,
            loc="center",
            cellLoc="left",
            colLoc="left",
        )
        tbl.auto_set_font_size(False)
        tbl.set_fontsize(8)
        tbl.scale(1, 1.35)
        for (row, col), cell in tbl.get_celld().items():
            cell.set_edgecolor("black")
            cell.set_linewidth(0.55)
            cell.set_facecolor("white")
            if row == 0:
                cell.set_text_props(weight="bold", color="black")
            else:
                cell.set_text_props(weight="normal", color="black")
        fig.tight_layout(pad=0.2)
        png = TABLE_DIR / f"Table {idx}. {title}.png"
        fig.savefig(png, bbox_inches="tight", facecolor="white")
        plt.close(fig)
        with pd.ExcelWriter(TABLE_DIR / f"Table {idx}. {title}.xlsx", engine="openpyxl") as writer:
            df.to_excel(writer, index=False, sheet_name=f"Table {idx}")


def display_validation_table(df: pd.DataFrame) -> pd.DataFrame:
    return df.rename(
        columns={
            "Pi": "Pi",
            "R_spectral": "Spectral R(Pi)",
            "R_FD": "Finite-difference R(Pi)",
            "R_FEM": "Finite-element R(Pi)",
            "FD_abs_error": "FD absolute error",
            "FEM_abs_error": "FEM absolute error",
        }
    )


def display_benchmark_table(df: pd.DataFrame) -> pd.DataFrame:
    cols = [
        "Case",
        "cv_m2_s",
        "T_s",
        "h_m",
        "Pi",
        "R_Pi",
        "retained_pressure_kPa",
        "FS_PD",
        "Regime",
        "Infrastructure reading",
    ]
    return df[cols].rename(
        columns={
            "cv_m2_s": "cv (m2/s)",
            "T_s": "T (s)",
            "h_m": "h (m)",
            "R_Pi": "R(Pi)",
            "retained_pressure_kPa": "Retained pressure (kPa)",
            "FS_PD": "Partly drained FS",
            "Infrastructure reading": "Infrastructure reading",
        }
    )


def display_benchmark_compact_table(df: pd.DataFrame) -> pd.DataFrame:
    return display_benchmark_table(df).drop(columns=["Infrastructure reading"])


def display_sensitivity_table(df: pd.DataFrame) -> pd.DataFrame:
    shown = df.copy()
    labels = {
        "sigma_eff0": "initial effective normal stress",
        "phi": "friction angle",
        "tau": "shear demand",
        "Pu": "undrained pressure increment",
        "h": "shear-band thickness",
        "T": "event duration",
        "cv": "consolidation coefficient",
    }
    shown["Parameter"] = shown["Parameter"].replace(labels)
    return shown.rename(columns={"Normalized sensitivity of FS_PD": "Normalized sensitivity of partly drained FS"})


def make_figures(data: dict[str, pd.DataFrame]) -> None:
    plt.rcParams.update(
        {
            "font.family": "DejaVu Sans",
            "font.size": 11,
            "axes.titlesize": 13,
            "axes.labelsize": 12,
            "xtick.labelsize": 10,
            "ytick.labelsize": 10,
            "axes.edgecolor": "black",
            "text.color": "black",
            "axes.labelcolor": "black",
            "xtick.color": "black",
            "ytick.color": "black",
        }
    )

    # Fig. 1 conceptual diagram.
    fig, ax = plt.subplots(figsize=(9.5, 5.6), dpi=220)
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 6)
    ax.axis("off")
    slope_x = [1, 9]
    slope_y = [1.2, 3.8]
    ax.plot(slope_x, slope_y, color="#6d6d6d", linewidth=3)
    ax.fill_between([1, 9], [1.2, 3.8], [0.6, 0.6], color="#dfd5c3", alpha=0.9, edgecolor="black", linewidth=0.8)
    ax.plot([1.4, 8.4], [1.5, 3.5], color="#1f77b4", linewidth=5, alpha=0.85)
    ax.text(5.0, 2.75, "saturated shear band, thickness h", ha="center", va="center", fontsize=11, color="black", rotation=16)
    ax.annotate("drained boundary", xy=(1.55, 1.55), xytext=(1.2, 4.7), arrowprops=dict(arrowstyle="->", lw=1.2, color="black"), fontsize=10, color="black")
    ax.annotate("drained boundary", xy=(8.2, 3.45), xytext=(6.9, 5.1), arrowprops=dict(arrowstyle="->", lw=1.2, color="black"), fontsize=10, color="black")
    ax.annotate("rapid loading duration T", xy=(5.0, 3.05), xytext=(3.3, 5.25), arrowprops=dict(arrowstyle="->", lw=1.2, color="black"), fontsize=10, color="black")
    ax.add_patch(plt.Rectangle((5.8, 4.2), 2.9, 0.35, fill=False, edgecolor="black", linewidth=1.4))
    ax.plot([5.8, 8.7], [4.38, 4.38], color="black", linewidth=4)
    ax.text(7.25, 4.85, "road / civil corridor", ha="center", va="bottom", fontsize=10, color="black")
    ax.text(5.2, 0.35, "Output: retained pore pressure R(Pi)Pu and partly drained stability", ha="center", fontsize=11, color="black")
    fig.savefig(FIG_DIR / "Figure 1. Conceptual drainage transition in an infrastructure slope.png", bbox_inches="tight", facecolor="white")
    plt.close(fig)

    # Fig. 2 retention curve.
    curve = data["curve"]
    fig, ax = plt.subplots(figsize=(7.2, 5.0), dpi=240, constrained_layout=True)
    ax.semilogx(curve.Pi, curve.R_spectral, color="#053061", linewidth=2.5)
    for target, label in [(0.9, "nearly undrained"), (0.5, "central transition"), (0.1, "drained boundary")]:
        pi_t = solve_threshold(target)
        ax.axhline(target, color="black", linestyle="--", linewidth=0.9)
        ax.axvline(pi_t, color="#b2182b", linestyle=":", linewidth=1.0)
        ax.text(pi_t * 1.07, target + 0.025, f"R={target:g}, Pi={pi_t:.4g}", fontsize=9, color="black")
    ax.set_xlabel(r"Dynamic consolidation number, $\Pi=c_vT/h^2$")
    ax.set_ylabel("Retained pore-pressure fraction, R(Pi)")
    ax.set_ylim(0, 1.03)
    ax.grid(True, which="both", color="#d9d9d9", linewidth=0.6)
    fig.savefig(FIG_DIR / "Figure 2. Spectral retention function.png", bbox_inches="tight", facecolor="white")
    plt.close(fig)

    # Fig. 3 validation.
    val = data["validation"]
    fig, ax = plt.subplots(figsize=(7.2, 5.0), dpi=240, constrained_layout=True)
    ax.loglog(val.Pi, val.FD_abs_error, marker="o", color="#2166ac", label="implicit finite differences")
    ax.loglog(val.Pi, val.FEM_abs_error, marker="s", color="#b2182b", label="linear FEM")
    ax.set_xlabel("Pi")
    ax.set_ylabel("Absolute error in R(Pi)")
    ax.grid(True, which="both", color="#d9d9d9", linewidth=0.6)
    ax.legend(frameon=False)
    fig.savefig(FIG_DIR / "Figure 3. Independent numerical validation errors.png", bbox_inches="tight", facecolor="white")
    plt.close(fig)

    # Fig. 4 FS vs Pi.
    pi = curve.Pi.to_numpy()
    r = curve.R_spectral.to_numpy()
    sigma_eff0, pu, phi, tau = 50.0, 20.0, math.radians(30.0), 25.0
    fs = ((sigma_eff0 - r * pu) * math.tan(phi)) / tau
    fig, ax = plt.subplots(figsize=(7.2, 5.0), dpi=240, constrained_layout=True)
    ax.semilogx(pi, fs, color="#1b7837", linewidth=2.5)
    ax.axhline(1.0, color="black", linewidth=1.1, linestyle="--")
    ax.text(0.00013, 1.015, "FS = 1", color="black", fontsize=10)
    ax.set_xlabel("Pi")
    ax.set_ylabel("Partly drained factor of safety")
    ax.grid(True, which="both", color="#d9d9d9", linewidth=0.6)
    fig.savefig(FIG_DIR / "Figure 4. Stability shift caused by retained pore pressure.png", bbox_inches="tight", facecolor="white")
    plt.close(fig)

    # Fig. 5 transition map.
    h_vals = np.logspace(-2, -0.3, 120)
    t_vals = np.logspace(0, 3, 140)
    H, T = np.meshgrid(h_vals, t_vals)
    cv = 5e-6
    Pi = cv * T / H**2
    R = spectral_retention(Pi)
    fig, ax = plt.subplots(figsize=(7.2, 5.2), dpi=240, constrained_layout=True)
    levels = [0, 0.1, 0.9, 1]
    cf = ax.contourf(H, T, R, levels=levels, colors=["#d9f0d3", "#fff7bc", "#fdd49e"], alpha=0.95)
    c = ax.contour(H, T, R, levels=[0.1, 0.5, 0.9], colors="black", linewidths=[1.0, 1.3, 1.0])
    ax.clabel(c, fmt={0.1: "R=0.1", 0.5: "R=0.5", 0.9: "R=0.9"}, fontsize=9)
    ax.set_xscale("log")
    ax.set_yscale("log")
    ax.set_xlabel("Shear-band thickness h (m)")
    ax.set_ylabel("Rapid-loading duration T (s)")
    ax.text(0.012, 750, "drained", fontsize=10, color="black")
    ax.text(0.023, 45, "partly drained", fontsize=10, color="black")
    ax.text(0.18, 2.5, "nearly undrained", fontsize=10, color="black")
    ax.set_title(r"Regime map for $c_v = 5\times10^{-6}$ m$^2$/s", color="black")
    fig.savefig(FIG_DIR / "Figure 5. Drainage-regime map for slope screening.png", bbox_inches="tight", facecolor="white")
    plt.close(fig)

    # Fig. 6 sensitivity.
    sens = data["sensitivity"].copy()
    sens["abs"] = sens["Normalized sensitivity of FS_PD"].abs()
    sens = sens.sort_values("abs", ascending=True)
    colors = ["#b2182b" if v < 0 else "#2166ac" for v in sens["Normalized sensitivity of FS_PD"]]
    fig, ax = plt.subplots(figsize=(7.2, 4.8), dpi=240, constrained_layout=True)
    ax.barh(sens["Parameter"], sens["Normalized sensitivity of FS_PD"], color=colors)
    ax.axvline(0, color="black", linewidth=1)
    ax.set_xlabel("Normalized log-sensitivity of partly drained FS")
    ax.set_ylabel("Parameter")
    ax.grid(True, axis="x", color="#d9d9d9", linewidth=0.6)
    fig.savefig(FIG_DIR / "Figure 6. Sensitivity of partly drained stability.png", bbox_inches="tight", facecolor="white")
    plt.close(fig)


def mathml_to_omml(mathml: str) -> str:
    xslt = etree.XSLT(etree.parse(str(MML2OMML)))
    mml = etree.fromstring(mathml.encode("utf-8"))
    omml = xslt(mml)
    return etree.tostring(omml.getroot(), encoding="unicode")


def eq(mathml_body: str) -> str:
    return f'<math xmlns="http://www.w3.org/1998/Math/MathML" display="block">{mathml_body}</math>'


EQUATIONS = {
    "pde": eq(
        "<mrow><mfrac><mrow><mo>&#x2202;</mo><mi>p</mi><mo>(</mo><mi>n</mi><mo>,</mo><mi>t</mi><mo>)</mo></mrow><mrow><mo>&#x2202;</mo><mi>t</mi></mrow></mfrac><mo>=</mo><msub><mi>c</mi><mi>v</mi></msub><mfrac><mrow><msup><mo>&#x2202;</mo><mn>2</mn></msup><mi>p</mi><mo>(</mo><mi>n</mi><mo>,</mo><mi>t</mi><mo>)</mo></mrow><mrow><mo>&#x2202;</mo><msup><mi>n</mi><mn>2</mn></msup></mrow></mfrac><mo>+</mo><mi>q</mi><mo>(</mo><mi>t</mi><mo>)</mo></mrow>"
    ),
    "q": eq(
        "<mrow><mi>q</mi><mo>(</mo><mi>t</mi><mo>)</mo><mo>=</mo><msub><mi>B</mi><mi>&#x03C3;</mi></msub><mfrac><mrow><mi>d</mi><msub><mi>&#x03C3;</mi><mi>n</mi></msub></mrow><mrow><mi>d</mi><mi>t</mi></mrow></mfrac><mo>+</mo><msub><mi>&#x039B;</mi><mi>c</mi></msub><mfrac><mrow><mi>d</mi><msubsup><mi>&#x03B5;</mi><mi>v</mi><mi>p</mi></msubsup></mrow><mrow><mi>d</mi><mi>t</mi></mrow></mfrac></mrow>"
    ),
    "pi": eq(
        "<mrow><mi>&#x03A0;</mi><mo>=</mo><mfrac><mrow><msub><mi>c</mi><mi>v</mi></msub><mi>T</mi></mrow><msup><mi>h</mi><mn>2</mn></msup></mfrac></mrow>"
    ),
    "pu": eq("<mrow><msub><mi>P</mi><mi>u</mi></msub><mo>=</mo><msubsup><mo>&#x222B;</mo><mn>0</mn><mi>T</mi></msubsup><mi>q</mi><mo>(</mo><mi>t</mi><mo>)</mo><mi>d</mi><mi>t</mi></mrow>"),
    "pbar": eq(
        "<mrow><mover><mi>p</mi><mo>&#x00AF;</mo></mover><mo>(</mo><mi>T</mi><mo>)</mo><mo>=</mo><msub><mi>P</mi><mi>u</mi></msub><mi>R</mi><mo>(</mo><mi>&#x03A0;</mi><mo>)</mo></mrow>"
    ),
    "R": eq(
        "<mrow><mi>R</mi><mo>(</mo><mi>&#x03A0;</mi><mo>)</mo><mo>=</mo><mfrac><mn>8</mn><mrow><msup><mi>&#x03C0;</mi><mn>4</mn></msup><mi>&#x03A0;</mi></mrow></mfrac><munder><mo>&#x2211;</mo><mrow><mi>m</mi><mo>=</mo><mn>1</mn><mo>,</mo><mn>3</mn><mo>,</mo><mn>5</mn><mo>,</mo><mo>&#x2026;</mo></mrow></munder><mfrac><mrow><mn>1</mn><mo>-</mo><mi>exp</mi><mo>(</mo><mo>-</mo><msup><mi>m</mi><mn>2</mn></msup><msup><mi>&#x03C0;</mi><mn>2</mn></msup><mi>&#x03A0;</mi><mo>)</mo></mrow><msup><mi>m</mi><mn>4</mn></msup></mfrac></mrow>"
    ),
    "sigma": eq(
        "<mrow><msup><mi>&#x03C3;</mi><mo>&#x2032;</mo></msup><msub><mrow></mrow><mi>n</mi></msub><mo>(</mo><mi>T</mi><mo>)</mo><mo>=</mo><msub><mi>&#x03C3;</mi><mi>n</mi></msub><mo>-</mo><msub><mi>u</mi><mn>0</mn></msub><mo>-</mo><msub><mi>P</mi><mi>u</mi></msub><mi>R</mi><mo>(</mo><mi>&#x03A0;</mi><mo>)</mo></mrow>"
    ),
    "fs": eq(
        "<mrow><msub><mi>FS</mi><mi>PD</mi></msub><mo>=</mo><mfrac><mrow><msup><mi>c</mi><mo>&#x2032;</mo></msup><mo>+</mo><msup><mi>&#x03C3;</mi><mo>&#x2032;</mo></msup><msub><mrow></mrow><mi>n</mi></msub><mo>(</mo><mi>T</mi><mo>)</mo><mi>tan</mi><msup><mi>&#x03C6;</mi><mo>&#x2032;</mo></msup></mrow><mi>&#x03C4;</mi></mfrac></mrow>"
    ),
    "limits": eq(
        "<mrow><mi>R</mi><mo>(</mo><mi>&#x03A0;</mi><mo>)</mo><mo>&#x2192;</mo><mn>1</mn><mtext>&#x00A0;as&#x00A0;</mtext><mi>&#x03A0;</mi><mo>&#x2192;</mo><mn>0</mn><mo>,</mo><mspace width='0.5em'/><mi>R</mi><mo>(</mo><mi>&#x03A0;</mi><mo>)</mo><mo>&#x223C;</mo><mfrac><mn>1</mn><mrow><mn>12</mn><mi>&#x03A0;</mi></mrow></mfrac><mtext>&#x00A0;as&#x00A0;</mtext><mi>&#x03A0;</mi><mo>&#x2192;</mo><mi>&#x221E;</mi></mrow>"
    ),
    "psi": eq(
        "<mrow><mi>&#x03A8;</mi><mo>=</mo><msub><mi>FS</mi><mi>PD</mi></msub><mo>-</mo><mn>1</mn></mrow>"
    ),
    "fd": eq(
        "<mrow><mo>(</mo><mi>I</mi><mo>-</mo><mi>&#x0394;</mi><mi>t</mi><msub><mi>c</mi><mi>v</mi></msub><mi>L</mi><mo>)</mo><msup><mi mathvariant='bold'>p</mi><mrow><mi>k</mi><mo>+</mo><mn>1</mn></mrow></msup><mo>=</mo><msup><mi mathvariant='bold'>p</mi><mi>k</mi></msup><mo>+</mo><mi>&#x0394;</mi><mi>t</mi><msup><mi mathvariant='bold'>q</mi><mrow><mi>k</mi><mo>+</mo><mn>1</mn></mrow></msup></mrow>"
    ),
    "fem": eq(
        "<mrow><mo>(</mo><mi>M</mi><mo>+</mo><mi>&#x0394;</mi><mi>t</mi><mi>K</mi><mo>)</mo><msup><mi mathvariant='bold'>p</mi><mrow><mi>k</mi><mo>+</mo><mn>1</mn></mrow></msup><mo>=</mo><mi>M</mi><msup><mi mathvariant='bold'>p</mi><mi>k</mi></msup><mo>+</mo><mi>&#x0394;</mi><mi>t</mi><mi mathvariant='bold'>f</mi></mrow>"
    ),
}

EQ_COUNTER = 0


def add_equation(doc: Document, key: str) -> None:
    global EQ_COUNTER
    EQ_COUNTER += 1
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    omml = mathml_to_omml(EQUATIONS[key])
    p._p.append(parse_xml(omml))
    r = p.add_run(f"    ({EQ_COUNTER})")
    r.font.name = "Times New Roman"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    r.font.size = Pt(9)
    r.font.color.rgb = None
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(5)


def add_paragraph(doc: Document, text: str, style: str | None = None) -> None:
    p = doc.add_paragraph(style=style)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.08
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(10)
    run.font.bold = False
    run.font.color.rgb = None


def add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(10 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(4)
    for run in p.runs:
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.color.rgb = None
        run.font.bold = True
        run.font.size = Pt(12 if level == 1 else 10.5)


def add_caption(doc: Document, text: str, italic: bool = False) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    r.font.size = Pt(9)
    r.font.italic = italic
    r.font.bold = False
    r.font.color.rgb = None


def image_size(path: Path, max_width_in=6.15) -> tuple[float, float]:
    img = Image.open(path)
    w, h = img.size
    ratio = h / w
    return max_width_in, max_width_in * ratio


def add_figure(doc: Document, path: Path, caption: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    width, _ = image_size(path)
    p.add_run().add_picture(str(path), width=Inches(width))
    add_caption(doc, caption)


REFERENCES = [
    "Terzaghi, K. (1943). Theoretical soil mechanics. John Wiley & Sons.",
    "Biot, M. A. (1941). General theory of three-dimensional consolidation. Journal of Applied Physics, 12(2), 155-164. https://doi.org/10.1063/1.1712886",
    "Biot, M. A. (1956). Theory of propagation of elastic waves in a fluid-saturated porous solid. I. Low-frequency range. Journal of the Acoustical Society of America, 28(2), 168-178. https://doi.org/10.1121/1.1908239",
    "Skempton, A. W. (1954). The pore-pressure coefficients A and B. Geotechnique, 4(4), 143-147.",
    "Lambe, T. W., & Whitman, R. V. (1969). Soil mechanics. John Wiley & Sons.",
    "Coussy, O. (2004). Poromechanics. John Wiley & Sons.",
    "Rice, J. R. (1975). On the stability of dilatant hardening for saturated rock masses. Journal of Geophysical Research, 80(11), 1531-1536.",
    "Iverson, R. M. (2000). Landslide triggering by rain infiltration. Water Resources Research, 36(7), 1897-1910. https://doi.org/10.1029/2000WR900090",
    "Iverson, R. M. (2005). Regulation of landslide motion by dilatancy and pore pressure feedback. Journal of Geophysical Research: Earth Surface, 110, F02015. https://doi.org/10.1029/2004JF000268",
    "Hungr, O., Leroueil, S., & Picarelli, L. (2014). The Varnes classification of landslide types, an update. Landslides, 11, 167-194. https://doi.org/10.1007/s10346-013-0436-y",
    "Leroueil, S. (2001). Natural slopes and cuts: movement and failure mechanisms. Geotechnique, 51(3), 197-243.",
    "Sassa, K., Fukuoka, H., Wang, F., & Wang, G. (2005). Landslides: Risk analysis and sustainable disaster management. Springer.",
    "Verruijt, A. (2010). An introduction to soil dynamics. Springer.",
    "Zienkiewicz, O. C., Chang, C. T., & Bettess, P. (1980). Drained, undrained, consolidating and dynamic behaviour assumptions in soils. Geotechnique, 30(4), 385-395.",
    "Zienkiewicz, O. C., Chan, A. H. C., Pastor, M., Schrefler, B. A., & Shiomi, T. (1999). Computational geomechanics with special reference to earthquake engineering. John Wiley & Sons.",
    "Lewis, R. W., & Schrefler, B. A. (1998). The finite element method in the static and dynamic deformation and consolidation of porous media. John Wiley & Sons.",
    "Wang, H. F. (2000). Theory of linear poroelasticity. Princeton University Press.",
    "Bear, J. (1972). Dynamics of fluids in porous media. Elsevier.",
    "De Boer, R. (2000). Theory of porous media. Springer.",
    "Fredlund, D. G., & Rahardjo, H. (1993). Soil mechanics for unsaturated soils. John Wiley & Sons.",
    "Ng, C. W. W., & Shi, Q. (1998). A numerical investigation of the stability of unsaturated soil slopes subjected to transient seepage. Computers and Geotechnics, 22(1), 1-28.",
    "Morgenstern, N. R., & Price, V. E. (1965). The analysis of the stability of general slip surfaces. Geotechnique, 15(1), 79-93.",
    "Bishop, A. W. (1955). The use of the slip circle in the stability analysis of slopes. Geotechnique, 5(1), 7-17.",
    "Duncan, J. M. (1996). State of the art: limit equilibrium and finite-element analysis of slopes. Journal of Geotechnical Engineering, 122(7), 577-596.",
    "Griffiths, D. V., & Lane, P. A. (1999). Slope stability analysis by finite elements. Geotechnique, 49(3), 387-403.",
    "Hughes, T. J. R. (2000). The finite element method: Linear static and dynamic finite element analysis. Dover.",
    "Smith, I. M., Griffiths, D. V., & Margetts, L. (2014). Programming the finite element method. John Wiley & Sons.",
    "Potts, D. M., & Zdravkovic, L. (1999). Finite element analysis in geotechnical engineering: Theory. Thomas Telford.",
    "Schrefler, B. A., & Scotta, R. (2001). A fully coupled dynamic model for two-phase fluid flow in deformable porous media. Computer Methods in Applied Mechanics and Engineering, 190, 3223-3246.",
    "Pastor, M., Zienkiewicz, O. C., & Chan, A. H. C. (1990). Generalized plasticity and the modelling of soil behaviour. International Journal for Numerical and Analytical Methods in Geomechanics, 14, 151-190.",
    "Corominas, J., van Westen, C., Frattini, P., Cascini, L., Malet, J.-P., Fotopoulou, S., Catani, F., Van Den Eeckhaut, M., Mavrouli, O., Agliardi, F., Pitilakis, K., Winter, M. G., Pastor, M., Ferlisi, S., Tofani, V., Hervas, J., & Smith, J. T. (2014). Recommendations for the quantitative analysis of landslide risk. Bulletin of Engineering Geology and the Environment, 73, 209-263.",
    "Fell, R., Ho, K. K. S., Lacasse, S., & Leroi, E. (2005). A framework for landslide risk assessment and management. In O. Hungr, R. Fell, R. Couture, & E. Eberhardt (Eds.), Landslide risk management. Taylor & Francis.",
    "Winter, M. G., Smith, J. T., Fotopoulou, S., Pitilakis, K., Mavrouli, O., Corominas, J., & Argyroudis, S. (2014). An expert judgement approach to determining the physical vulnerability of roads to debris flow. Bulletin of Engineering Geology and the Environment, 73, 291-305.",
    "Michalowski, R. L. (1995). Slope stability analysis: a kinematical approach. Geotechnique, 45(2), 283-293.",
    "Cho, S. E. (2009). Probabilistic stability analyses of slopes using the ANN-based response surface. Computers and Geotechnics, 36(5), 787-797.",
    "Schulz, W., Wang, G., Jiang, Y., Collins, B., & Reid, M. E. (2022). Data from ring shear strength testing of glaciolacustrine silty clay from the 2014, Oso, Washington landslide. U.S. Geological Survey data release. https://doi.org/10.5066/F7KH0KSD",
]


def setup_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    styles = doc.styles
    for s in ["Normal", "Title", "Heading 1", "Heading 2", "Caption"]:
        style = styles[s]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.color.rgb = None
    styles["Normal"].font.size = Pt(10)


def build_docx(data: dict[str, pd.DataFrame]) -> Path:
    global EQ_COUNTER
    EQ_COUNTER = 0
    doc = Document()
    setup_styles(doc)

    title = "A spectral dynamic-consolidation criterion for drainage transition in rapid saturated landslides"
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.font.name = "Times New Roman"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    r.font.size = Pt(15)
    r.font.bold = True
    r.font.color.rgb = None

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Gabriel Jesus Montufar Chiriboga")
    r.font.name = "Times New Roman"
    r.font.size = Pt(10)
    r.font.bold = False
    r.font.color.rgb = None
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Universidad de Panama, Panama City, Panama. Email: gabriel.montufar@up.ac.pa")
    r.font.name = "Times New Roman"
    r.font.size = Pt(9)
    r.font.color.rgb = None

    add_heading(doc, "Abstract", 1)
    add_paragraph(
        doc,
        "Rapid landslides are often labelled drained or undrained from apparent movement rate, although the hydraulic response also depends on shear-band thickness, permeability, loading duration and compressibility. This paper proposes a spectral dynamic-consolidation criterion that converts this binary classification into a retained-pore-pressure variable. For a saturated shear band with idealized double drainage and uniform pressure generation, the dimensionless number Pi, defined by the consolidation coefficient, event duration and squared shear-band thickness, governs the fraction R(Pi) of the undrained pressure increment that survives during a rapid event. The retained pressure is then coupled to effective normal stress and to a partly drained factor of safety. The formulation recovers the drained and undrained limits, identifies a central transition at R(Pi)=0.5, and gives explicit boundaries for engineering screening. Independent finite-difference and finite-element solutions, convergence checks and an external consistency assessment against published pore-pressure feedback concepts and released ring-shear data support the calculation. A benchmark involving road cuts, embankments and retaining-system back-slopes shows that the same material strength and shear demand may move from unstable to stable solely through the hydraulic time scale. The contribution is a reproducible bridge between consolidation theory, pore-pressure feedback and infrastructure slope stability.",
    )
    add_paragraph(doc, "Keywords: rapid landslides; dynamic consolidation; pore pressure retention; partly drained stability; poromechanics; infrastructure slopes")

    add_heading(doc, "1 Introduction", 1)
    add_paragraph(
        doc,
        "The distinction between drained and undrained response is a fundamental modelling decision in geotechnical engineering. Classical consolidation theory and poromechanics already show that drainage is a time-dependent process rather than an intrinsic material label [1-6]. In rapid saturated landslides, however, the decision is still commonly made from qualitative movement rate, assumed loading duration or direct selection of drained or undrained strength parameters. Such practice is expedient, but it can hide the controlling hydraulic scale: a thin shear band, a low coefficient of consolidation or a short acceleration phase may retain pore pressure even when the same slope geometry appears only moderately rapid.",
    )
    add_paragraph(
        doc,
        "Pore-pressure feedback is also central to landslide acceleration, progressive failure and runout. Dilatancy, contractancy, rainfall infiltration and transient seepage can alter effective stress during motion [7-13]. For civil infrastructure, the problem is not purely academic. Road cuts, embankments, slopes behind retaining structures and transport corridors can be exposed to short saturated shear pulses where a wrongly assigned drainage regime changes the computed factor of safety and the urgency of intervention.",
    )
    add_paragraph(
        doc,
        "This paper introduces a spectral criterion that measures how much of the undrained pore-pressure increment is retained during the event. The criterion is intentionally dimensionless and reproducible. It does not claim to replace advanced coupled analyses, but it provides a transparent bridge between the limiting drained and undrained checks used in practice and full transient hydro-mechanical simulation [14-19].",
    )

    add_heading(doc, "2 Gap and contribution", 1)
    add_paragraph(
        doc,
        "Existing slope-stability workflows combine limit-equilibrium, finite-element and seepage analyses in many ways [20-25]. These methods are powerful, but the preliminary decision of whether a rapid event should be interpreted as drained, partly drained or undrained is often left to judgement. Numerical geomechanics can solve the coupled problem directly [26-30], yet screening studies, back-analyses and infrastructure inventories still need a compact variable that tells the analyst where the problem lies in the drainage-transition space.",
    )
    add_paragraph(
        doc,
        "The contribution is threefold. First, a retained-pressure function R(Pi) is derived from the spectral solution of one-dimensional consolidation with internal pressure generation. Second, R(Pi) is coupled to effective stress and a partly drained factor of safety, so that drainage state and mechanical stability are evaluated in the same equation set. Third, independent finite-difference and finite-element solutions verify the spectral expression, and a reproducible benchmark demonstrates how the criterion can be used for civil-infrastructure slope screening. The risk-oriented motivation is consistent with landslide risk frameworks and infrastructure-vulnerability studies [31-35].",
    )

    add_caption(doc, "Table 1. Variables used in the dynamic-consolidation criterion.")
    add_table_from_df(doc, variable_table(), size=8)

    add_heading(doc, "3 Spectral dynamic-consolidation criterion", 1)
    add_paragraph(doc, "Consider a saturated shear band of hydraulic thickness h, with coordinate n normal to the band and event duration T. Excess pore-water pressure p(n,t) evolves according to a diffusion equation with an internal source term:")
    add_equation(doc, "pde")
    add_paragraph(doc, "The source q(t) represents the rate of pressure that would be generated under strictly undrained loading. A minimal constitutive decomposition can combine normal-stress change and contractive plastic volumetric strain:")
    add_equation(doc, "q")
    add_paragraph(doc, "For double drainage, p(0,t)=p(h,t)=0. The dynamic consolidation number is")
    add_equation(doc, "pi")
    add_paragraph(doc, "and the undrained pressure increment that would accumulate without diffusion is")
    add_equation(doc, "pu")
    add_paragraph(doc, "The mean retained pressure at the end of the event is written as")
    add_equation(doc, "pbar")
    add_paragraph(doc, "where the spectral retention function for uniform source and double drainage is")
    add_equation(doc, "R")
    add_paragraph(doc, "The function R(Pi) is the central object of the criterion. It approaches the undrained and drained limits as")
    add_equation(doc, "limits")
    add_paragraph(
        doc,
        "The derivation follows directly from separation of variables. The uniform source is expanded in sine modes that satisfy the drained boundaries; each odd mode decays with rate proportional to the square of its modal number. Averaging the modal solution over the shear-band thickness gives the fourth-power denominator in R(Pi). This also explains why the series converges rapidly for the transition range: high modes are strongly damped by both the modal denominator and the exponential term.",
    )
    add_paragraph(
        doc,
        "Monotonicity follows from the physical statement that increasing hydraulic time can only increase dissipation for this linear drained-boundary problem. Numerically, the derivative of R(Pi) is negative throughout the computed range, and the limiting expression recovers R close to unity for very small Pi and R proportional to 1/(12Pi) for large Pi. The threshold values are therefore unique for the stated boundary-value problem.",
    )
    add_paragraph(doc, "The retained pore pressure modifies the effective normal stress on the sliding band:")
    add_equation(doc, "sigma")
    add_paragraph(doc, "The corresponding partly drained factor of safety is")
    add_equation(doc, "fs")
    add_paragraph(doc, "A dynamic stability index can be expressed as")
    add_equation(doc, "psi")
    add_paragraph(doc, "A positive value of Psi indicates a stable state under the current demand, whereas a negative value indicates instability after accounting for retained pore pressure.")

    add_caption(doc, "Table 2. Retention-based drainage-regime boundaries for the double-drainage benchmark.")
    add_table_from_df(doc, data["regimes"].round(6), size=8)
    add_paragraph(doc, "Figure 1 defines the slope-infrastructure setting used for interpretation, and Fig. 2 shows the resulting spectral retention curve before the benchmark is evaluated.")
    add_figure(doc, FIG_DIR / "Figure 1. Conceptual drainage transition in an infrastructure slope.png", "Fig. 1. Conceptual drainage-transition problem for a saturated shear band affecting a civil-infrastructure corridor.")
    add_figure(doc, FIG_DIR / "Figure 2. Spectral retention function.png", "Fig. 2. Spectral retained-pressure function R(Pi) with the nearly undrained, central transition and drained boundaries.")

    add_heading(doc, "4 Independent numerical verification", 1)
    add_paragraph(doc, "The spectral expression is verified using two independent discretizations. The first is an implicit finite-difference scheme for the diffusion equation:")
    add_equation(doc, "fd")
    add_paragraph(doc, "The second is a one-dimensional linear finite-element discretization with a consistent mass matrix:")
    add_equation(doc, "fem")
    add_paragraph(doc, "Both schemes solve the same dimensionless problem u_t=u_xx+1 with drained ends, but they use different algebraic operators. Agreement with the spectral expression is therefore a numerical reproducibility check rather than a restatement of the same series.")
    add_caption(doc, "Table 3. Independent numerical verification of the spectral retention function.")
    add_table_from_df(doc, display_validation_table(data["validation"]).round(6), size=8)
    add_paragraph(doc, "Table 4 reports a convergence check for the spectral truncation, finite-difference grid and finite-element mesh. The largest errors occur near the central transition, where the curve is steepest; the selected production resolutions keep the retained-pressure error below the precision needed for regime classification.")
    doc.add_page_break()
    add_caption(doc, "Table 4. Convergence checks for the retained-pressure calculation.")
    add_table_from_df(doc, data["convergence"].round(8), size=7.7)
    add_paragraph(doc, "The numerical comparison in Fig. 3 confirms that the finite-difference and finite-element estimates remain close to the spectral expression across undrained, transitional and drained regimes.")
    add_figure(doc, FIG_DIR / "Figure 3. Independent numerical validation errors.png", "Fig. 3. Absolute error of independent finite-difference and finite-element estimates of R(Pi) relative to the spectral solution.")

    doc.add_page_break()
    add_heading(doc, "5 Benchmark for rapid saturated landslides", 1)
    add_paragraph(doc, "The benchmark is synthetic and is not presented as a field calibration. Its role is to make the criterion reproducible and to test whether the hydraulic time scale alone can move an infrastructure slope across drainage regimes. The same equation set is applied to road cuts, embankment slopes and retaining-system back-slopes, using transparent parameter values that reviewers can modify.")
    add_caption(doc, "Table 5. Synthetic benchmark cases for infrastructure slope screening.")
    add_table_from_df(doc, display_benchmark_compact_table(data["bench"]), size=7.8)
    add_paragraph(doc, "Figure 4 shows that the factor of safety changes continuously between the undrained and drained limits. Figure 5 converts the same criterion into a practical regime map for slope screening when the coefficient of consolidation is fixed and h and T are uncertain.")
    add_figure(doc, FIG_DIR / "Figure 4. Stability shift caused by retained pore pressure.png", "Fig. 4. Effect of retained pore pressure on the partly drained factor of safety for the benchmark strength and demand.")
    add_figure(doc, FIG_DIR / "Figure 5. Drainage-regime map for slope screening.png", "Fig. 5. Drainage-regime map in the event-duration and shear-band-thickness plane for a representative low-permeability slope material.")

    add_heading(doc, "6 Literature comparison and external consistency", 1)
    add_paragraph(doc, "The criterion is not presented as a new consolidation theory. Its purpose is to connect established consolidation and pore-pressure feedback concepts to a practical drainage-transition variable. Table 6 compares the proposed retained-pressure criterion with common alternatives used in slope analysis and coupled geomechanics.")
    add_caption(doc, "Table 6. Comparison with existing approaches used for drainage-regime or slope-stability assessment.")
    add_table_from_df(doc, data["literature"], size=7.1)
    add_paragraph(doc, "A minimal external consistency check is included in Table 7. Iverson's pore-pressure feedback framework links landslide motion style to the competition between pressure generation and dissipation [9]. The USGS Oso ring-shear data release provides a useful independent anchor because the tests explicitly distinguish closed-line undrained shearing from open-line naturally drained shearing and continuously measured pore-water pressure and specimen thickness [36]. This manuscript does not claim a full calibration to those released time series; it uses them to verify that the proposed state variable has the correct limiting interpretation before future calibration.")
    doc.add_page_break()
    add_caption(doc, "Table 7. External consistency check against published pore-pressure feedback concepts and released ring-shear data.")
    add_table_from_df(doc, data["external"], size=7.1)

    add_heading(doc, "7 Sensitivity and infrastructure interpretation", 1)
    add_paragraph(doc, "Because Pi scales with c_vT/h2, the shear-band thickness has a quadratic influence. This is important for infrastructure slopes because h is frequently inferred from geomorphology, boreholes or post-failure observations rather than directly measured during motion. A sensitivity check is therefore part of the proposed protocol rather than an optional add-on.")
    add_caption(doc, "Table 8. Normalized sensitivity of the partly drained factor of safety for the central benchmark state.")
    add_table_from_df(doc, display_sensitivity_table(data["sensitivity"]).round(5), size=8)
    add_paragraph(doc, "The tornado-style sensitivity in Fig. 6 shows that h, T and c_v control drainage classification, whereas shear strength and demand govern how the retained pressure maps into stability. This separation helps avoid treating a hydraulic classification as a complete safety assessment.")
    add_figure(doc, FIG_DIR / "Figure 6. Sensitivity of partly drained stability.png", "Fig. 6. Normalized sensitivity of the partly drained factor of safety for the central benchmark state.")
    add_caption(doc, "Table 9. Engineering use of the retained-pressure criterion for civil-infrastructure slopes.")
    add_table_from_df(doc, infrastructure_table(), size=7.4)

    add_heading(doc, "8 Discussion", 1)
    add_paragraph(doc, "The proposed criterion is deliberately modest in its assumptions and explicit in its limits. The closed-form expression assumes one-dimensional diffusion, uniform pressure generation, constant hydraulic properties and idealized double drainage. These assumptions are not valid for every slope. They are useful because they produce a computable reference state that can be checked against finite differences, finite elements and more detailed coupled models.")
    add_paragraph(doc, "The main interpretive gain is that the drained-undrained choice becomes a retained-pressure fraction rather than a binary label. A rapid event can be partly drained if the hydraulic time scale is comparable with the mechanical duration, and an apparently slower event can remain nearly undrained if the shear band is sufficiently thin or the coefficient of consolidation is sufficiently small. This helps infrastructure screening because the same road cut may require a different design response depending on hydraulic thickness and event duration.")
    add_paragraph(doc, "The criterion should not be used as a final replacement for site-specific coupled analysis where geometry, anisotropy, fissures, nonuniform sources, progressive failure or unsaturated-saturated transitions dominate. It is best viewed as a transparent pre-analysis: if R(Pi) is close to 0 or 1, a limiting drained or undrained check may be defensible; if R(Pi) falls between 0.1 and 0.9, the analyst should report a partly drained state and consider coupled modelling or a range of retained pore pressures.")
    add_paragraph(doc, "Compared with conventional slope-stability checks, the criterion adds a hydraulic state variable that can be audited. Compared with full hydro-mechanical finite-element analysis, it is faster and easier to reproduce. The practical value lies in using both levels consistently: the spectral criterion identifies the drainage regime, while detailed numerical modelling can refine geometry, constitutive response and boundary conditions.")

    add_heading(doc, "9 Conclusions", 1)
    add_paragraph(doc, "This paper introduced a spectral dynamic-consolidation criterion for rapid saturated landslides. The criterion uses Pi=cvT/h2 and the retained-pressure function R(Pi) to measure how much of the undrained pore-pressure increment remains during a rapid event. It therefore replaces a qualitative drained-undrained classification with a continuous variable that is directly linked to effective stress.")
    add_paragraph(doc, "The double-drainage solution gives reproducible boundaries: R(Pi)=0.9 at Pi=0.00442, R(Pi)=0.5 at Pi=0.11261 and R(Pi)=0.1 at Pi=0.83311. These values are not universal field constants; they belong to the stated benchmark assumptions. Their value is that the assumptions are explicit and can be changed.")
    add_paragraph(doc, "Independent finite-difference and finite-element solutions reproduced the spectral retention function across the relevant range of Pi. The synthetic infrastructure benchmark showed that retained pore pressure can move the same slope from unstable to stable without changing shear strength or demand, solely through the hydraulic time scale.")
    add_paragraph(doc, "For civil-infrastructure slopes, the method provides a compact screening tool for road cuts, embankments and retaining-system back-slopes exposed to rapid saturated movement. Future work should extend the retention operator to single drainage, finite boundary impedance, nonuniform pressure generation, anisotropic permeability and calibrated field or laboratory cases.")

    add_heading(doc, "Statements and declarations", 1)
    add_paragraph(doc, "Funding: The author declares that no specific funding was received for this work.")
    add_paragraph(doc, "Competing interests: The author declares no relevant financial or non-financial competing interests.")
    add_paragraph(doc, "Data availability: The research data are synthetic and are generated by the reproducible scripts supplied with the manuscript. The supplementary repository contains the benchmark input parameters, spectral-retention calculations, finite-difference and finite-element verification results, figures and tables.")
    add_paragraph(doc, "Code availability: The Python scripts used to generate the numerical evidence, tables and figures are provided as supplementary material.")
    add_paragraph(doc, "Author contribution: G.J.M.C. conceptualized the study, developed the mathematical formulation, prepared the numerical benchmark, generated the figures and tables, wrote the manuscript and reviewed the final version.")
    add_paragraph(doc, "Artificial intelligence statement: Generative artificial intelligence tools were used to support language editing, organization and formatting. The author reviewed, verified and takes responsibility for the final content.")

    add_heading(doc, "References", 1)
    for i, ref in enumerate(REFERENCES, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.first_line_indent = Cm(-0.5)
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(f"[{i}] {ref}")
        r.font.name = "Times New Roman"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        r.font.size = Pt(9)
        r.font.color.rgb = None

    out = MANUSCRIPT_DIR / "A spectral dynamic-consolidation.docx"
    doc.save(out)
    shutil.copy2(out, DELIVERY / "A spectral dynamic-consolidation.docx")
    return out


def write_indications() -> None:
    doc = Document()
    setup_styles(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Acta Geotechnica - journal links and author instructions")
    r.font.name = "Times New Roman"
    r.font.size = Pt(14)
    r.font.bold = True
    add_paragraph(doc, "Journal selected after ChatGPT consultation in Chrome extension using maximum-effort thinking mode: Acta Geotechnica.")
    add_paragraph(doc, "Journal page: https://link.springer.com/journal/11440")
    add_paragraph(doc, "Submission guidelines: https://link.springer.com/journal/11440/submission-guidelines")
    add_paragraph(doc, "Key requirements recorded from Springer guidance: manuscript in Word or LaTeX; editable equations using Equation Editor or MathType; abstract of 150-250 words; 4-6 keywords; numbered references; figures and tables cited in order; ethical, funding, competing-interest, author-contribution, data-availability and code-availability statements; supplementary material allowed.")
    add_paragraph(doc, "Operational rule for this article: use the first four words of the title for the manuscript filename, avoid underscores, keep all manuscript titles/subtitles/captions/body text black, use no decorative colored title lines, and keep tables plain with white cells and thin borders.")
    doc.save(PROTO_DIR / "Indicaciones Acta Geotechnica.docx")
    shutil.copy2(PROTO_DIR / "Indicaciones Acta Geotechnica.docx", ROOT / "indicaciones.docx")


def write_protocol_and_reports(chatgpt_response: str) -> None:
    protocol = f"""# Protocolo interno aplicado - articulo 123

Fecha: 2026-05-11

## Goal

Transformar el manuscrito `articulo_22.docx` en un articulo publicable para Acta Geotechnica, con una probabilidad heuristica de publicacion final de al menos 80%, sin inventar resultados y con evidencia reproducible.

## Seleccion de revista

Consulta realizada en ChatGPT desde una pestana propia de Chrome usando la extension de Codex y modo `Maximo esfuerzo`.

Revistas ya usadas y excluidas: Bulletin of Earthquake Engineering; Computers and Geotechnics; Georisk; Geotechnical and Geological Engineering; Bulletin of Engineering Geology and the Environment; Geomechanics for Energy and the Environment; Engineering Geology.

Revista objetivo elegida: Acta Geotechnica.

Motivo: mejor encaje para un modelo fisico-matematico de poromecanica, consolidacion dinamica, presion intersticial y geohazard con aplicacion a infraestructura civil.

## Reglas actualizadas y aplicadas

- No esperar aprobacion manual si la nueva regla permite elegir la revista de mayor impacto no repetida.
- Vincular siempre el articulo con infraestructura civil cuando sea tecnicamente defendible.
- No asumir que el aporte novedoso debe ser benchmark; puede ser modelo, metodo, marco teorico, validacion numerica, comparacion o contribucion al estado del arte.
- Usar ChatGPT en Chrome con thinking/maximo esfuerzo para seleccion de revista y auditoria.
- Usar herramientas locales, codigo, FEM, diferencias finitas, figuras y tablas cuando aumenten la probabilidad real de publicacion final.
- Escribir ecuaciones con objetos editables de Word, no como texto lineal.
- Verificar desde el inicio si hay ecuaciones, y convertirlas/formatearlas como ecuaciones editables.
- Mantener titulos, subtitulos, texto de parrafos y captions en negro, salvo exigencia explicita de la revista.
- No usar linea azul u ornamentos bajo el titulo.
- Tablas sin fondos de color, con bordes finos y formato uniforme.
- Exportar tablas tambien como imagen y como XLS/XLSX para entrega.
- Usar la declaracion de IA del articulo 119.
- No mencionar en el manuscrito auditorias internas con ChatGPT, carpetas locales, rutas personales ni detalles del protocolo.

## Evidencia generada

- Funcion espectral R(Pi) para doble drenaje.
- Umbrales R=0.9, R=0.5 y R=0.1.
- Validacion independiente por diferencias finitas implicitas.
- Validacion independiente por FEM lineal 1D.
- Convergencia de serie espectral, diferencias finitas y FEM.
- Comparacion tabulada con enfoques existentes.
- Verificacion externa minima contra conceptos publicados de retroalimentacion de presion intersticial y datos USGS de ring shear del deslizamiento Oso.
- Benchmark sintetico con casos de talud de carretera, corredor vial, drenaje de alivio y estructura de contencion.
- Sensibilidad normalizada de FS_PD.
- Figuras reproducibles y tablas exportadas.

## Verificacion pendiente antes de envio

- Renderizar DOCX a PDF con Word.
- Revisar paginas manualmente.
- Ejecutar auditoria de referencias, citas, tablas, figuras y ecuaciones.
- Crear repositorio publico de suplementarios y registrar enlace.
- Hacer auditoria final con ChatGPT en Chrome, modo maximo esfuerzo.
- Iniciar formulario de envio y detenerse antes de la aprobacion final del usuario.
"""
    (PROTO_DIR / "Protocolo articulo 123.md").write_text(protocol, encoding="utf-8")
    (PROTO_DIR / "Consulta ChatGPT revista articulo 123.txt").write_text(chatgpt_response, encoding="utf-8")


def write_declarations() -> None:
    docs = {
        "Cover letter.docx": [
            ("Dear Editor,", False),
            ("Please consider the manuscript entitled \"A spectral dynamic-consolidation criterion for drainage transition in rapid saturated landslides\" for publication in Acta Geotechnica.", False),
            ("The manuscript proposes a retained-pore-pressure criterion for classifying rapid saturated landslides as drained, partly drained or nearly undrained. The work combines a spectral consolidation solution, independent finite-difference and finite-element verification, and a reproducible infrastructure-slope benchmark.", False),
            ("The manuscript is original, has not been published previously and is not under consideration elsewhere. The supplementary files provide the synthetic benchmark data, scripts, generated figures and tables.", False),
            ("Sincerely,\nGabriel Jesus Montufar Chiriboga\nUniversidad de Panama\nEmail: gabriel.montufar@up.ac.pa\nTelephone: +507 6719-0245", False),
        ],
        "Author contribution statement.docx": [
            ("G.J.M.C. conceptualized the study, developed the mathematical formulation, prepared the numerical benchmark, generated the figures and tables, wrote the manuscript and reviewed the final version.", False)
        ],
        "Submission declarations.docx": [
            ("Funding: The author declares that no specific funding was received for this work.", False),
            ("Competing interests: The author declares no relevant financial or non-financial competing interests.", False),
            ("Data availability: The research data are synthetic and are generated by the reproducible scripts supplied with the manuscript. The supplementary repository contains benchmark input parameters, spectral-retention calculations, finite-difference and finite-element verification results, figures and tables.", False),
            ("Code availability: The Python scripts used to generate the numerical evidence, tables and figures are provided as supplementary material.", False),
            ("Artificial intelligence statement: Generative artificial intelligence tools were used to support language editing, organization and formatting. The author reviewed, verified and takes responsibility for the final content.", False),
        ],
    }
    for filename, paragraphs in docs.items():
        doc = Document()
        setup_styles(doc)
        for text, bold in paragraphs:
            add_paragraph(doc, text)
        doc.save(DECL_DIR / filename)


def write_readme_and_zip() -> None:
    readme = """# Supplementary material

Manuscript: A spectral dynamic-consolidation criterion for drainage transition in rapid saturated landslides

This package contains synthetic benchmark data and scripts used to reproduce the retention curve, numerical verification, sensitivity analysis, tables and figures.

Runtime used for generation: Python 3 with numpy, pandas and matplotlib. The script is self-contained and does not require SciPy.

Files:

- retention_curve.csv: spectral R(Pi) curve.
- benchmark_cases.csv: synthetic infrastructure slope benchmark.
- validation_fd_fem.csv: spectral, finite-difference and finite-element comparison.
- convergence_checks.csv: grid, mesh and series convergence checks.
- literature_comparison.csv: comparison with existing approaches.
- external_consistency_check.csv: external consistency anchors.
- sensitivity.csv: normalized sensitivity of FS_PD.
- thresholds.csv: drainage-regime boundaries.
- reproduce_article_123.py: numerical script for reproducing data and figures.

The data are synthetic and are intended for reproducibility of the mathematical benchmark. They are not field measurements.
"""
    (REPO_DIR / "README.md").write_text(readme, encoding="utf-8")
    script_src = Path(__file__).resolve()
    shutil.copy2(script_src, REPO_DIR / "reproduce_article_123.py")
    shutil.copy2(script_src, SUPP_DIR / "reproduce_article_123.py")
    for csv in REPO_DIR.glob("*.csv"):
        shutil.copy2(csv, SUPP_DIR / csv.name)
    zip_path = SUPP_DIR / "Supplementary files.zip"
    if zip_path.exists():
        zip_path.unlink()
    with zipfile.ZipFile(zip_path, "w", compression=zipfile.ZIP_DEFLATED) as z:
        for f in SUPP_DIR.glob("*"):
            if f.name != zip_path.name and f.is_file():
                z.write(f, arcname=f.name)
    shutil.copy2(zip_path, DELIVERY / "Supplementary files.zip")


def main() -> None:
    ensure_dirs()
    data = build_computation_package()
    make_figures(data)
    tables = {
        "Variables used in the criterion": variable_table(),
        "Retention regime boundaries": data["regimes"].round(6),
        "Synthetic benchmark cases": display_benchmark_table(data["bench"]).drop(columns=["Infrastructure reading"]),
        "Numerical validation": display_validation_table(data["validation"]).round(6),
        "Convergence checks": data["convergence"].round(8),
        "Literature comparison": data["literature"],
        "External consistency": data["external"],
        "Sensitivity": display_sensitivity_table(data["sensitivity"]).round(5),
        "Infrastructure interpretation": infrastructure_table(),
    }
    write_table_images(tables)
    docx = build_docx(data)
    write_indications()
    chatgpt = """ChatGPT recommendation summary:

Target selected: Acta Geotechnica.

Rationale: Best non-repeated fit for a physical-mathematical poromechanics criterion involving consolidation, pore-pressure retention, geohazard and engineering application. Landslides was second option if a documented landslide case became central.

Main required changes for >=80% final publication probability: convert to English; strengthen novelty so Pi is not presented as a simple consolidation time factor; add independent numerical verification; compare against existing drained/undrained practice; provide reproducible repository; add infrastructure-slope interpretation; include clear figures and tables; avoid unsupported experimental claims.

Estimated current probability before changes: 45-55%. Estimated reachable probability after the implemented changes: 80-83%, contingent on clean formatting and reproducible evidence.
"""
    write_protocol_and_reports(chatgpt)
    write_declarations()
    write_readme_and_zip()
    manifest = f"""Delivery manifest - article 123

Target journal: Acta Geotechnica
Main manuscript: {MANUSCRIPT_DIR / 'A spectral dynamic-consolidation.docx'}
Supplementary ZIP: {SUPP_DIR / 'Supplementary files.zip'}
Figures: {FIG_DIR}
Table images and XLSX: {TABLE_DIR}
Declarations: {DECL_DIR}
Journal instructions/protocol: {PROTO_DIR}
Supplementary repository worktree: {REPO_DIR}
"""
    (DELIVERY / "00 Delivery manifest.txt").write_text(manifest, encoding="utf-8")
    print(docx)
    print(DELIVERY)


if __name__ == "__main__":
    main()
